"""
Script: generate_instrumen_wawancara.py
Tujuan: Membuat file instrumen wawancara analisis kebutuhan SmartPresensi
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ──────────────────────────────────────────────
# Helper: set margins
# ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin   = Cm(4)
section.right_margin  = Cm(3)

# ──────────────────────────────────────────────
# Helper: warna & style
# ──────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1A, 0x3A, 0x6B)   # biru tua
BLUE_MID    = RGBColor(0x25, 0x5F, 0xA8)   # biru sedang (header tabel)
BLUE_LIGHT  = RGBColor(0xD6, 0xE4, 0xF7)   # biru muda (alt row)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)
GOLD        = RGBColor(0xD4, 0xA0, 0x17)

def set_cell_bg(cell, color: RGBColor):
    """Isi background warna sel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(color.red, color.green, color.blue)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set border sel tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:{}'.format(edge)
        element = OxmlElement(tag)
        element.set(qn('w:val'), kwargs.get('val', 'single'))
        element.set(qn('w:sz'), kwargs.get('sz', '6'))
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), kwargs.get('color', '25 5FA8'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def heading_para(doc, text, level=1, color=BLUE_DARK, size=14, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def normal_para(doc, text, size=11, bold=False, italic=False, color=None, space_before=0, space_after=4, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.alignment = align
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_divider(doc, color='1A3A6B'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════════════════
# HALAMAN JUDUL
# ══════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("INSTRUMEN WAWANCARA")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("ANALISIS KEBUTUHAN SISTEM INFORMASI")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run("PEMANTAUAN KEHADIRAN SISWA")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('"SmartPresensi"')
run.bold = True
run.italic = True
run.font.size = Pt(15)
run.font.color.rgb = BLUE_MID

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(20)
run = p.add_run("di SMKN 7 Purworejo")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = BLUE_DARK

add_divider(doc, '255F A8')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("SKRIPSI")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("Program Studi Pendidikan Teknik Informatika dan Komputer")
run.font.size = Pt(11)
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("Universitas Negeri Semarang")
run.font.size = Pt(11)
run.font.color.rgb = BLUE_DARK

add_divider(doc, '255F A8')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run("Aktor yang Diwawancarai:")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = BLUE_DARK

for aktor in ["1. Guru / Wali Kelas", "2. Siswa", "3. Orang Tua / Wali Murid"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(aktor)
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PETUNJUK UMUM
# ══════════════════════════════════════════════════════════

heading_para(doc, "A. PETUNJUK UMUM WAWANCARA", level=1, size=13)
add_divider(doc)

petunjuk = [
    "1. Wawancara ini bersifat semi-terstruktur dan open-ended, sehingga pewawancara dapat mengembangkan pertanyaan lanjutan (probing) sesuai respons narasumber.",
    "2. Jawaban narasumber dicatat secara ringkas pada kolom yang tersedia, dan dapat direkam dengan persetujuan narasumber.",
    "3. Tidak ada jawaban benar atau salah; setiap pendapat dan pengalaman narasumber sangat berharga bagi pengembangan sistem.",
    "4. Kerahasiaan identitas narasumber dijaga sesuai etika penelitian.",
    "5. Estimasi durasi wawancara: 20–35 menit per aktor.",
]
for item in petunjuk:
    normal_para(doc, item, size=11)

# ──────────────────────────────────────────────────────────
# Blueprint Dimensi Kebutuhan
# ──────────────────────────────────────────────────────────
heading_para(doc, "B. BLUEPRINT DIMENSI KEBUTUHAN (PIECES Framework + UX)", level=1, size=13, space_before=14)
add_divider(doc)

normal_para(doc,
    "Instrumen ini dirancang berdasarkan empat dimensi analisis kebutuhan yang disesuaikan "
    "dengan konteks sistem SmartPresensi:", size=11, space_after=8)

tbl_bp = doc.add_table(rows=5, cols=3)
tbl_bp.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_bp.style = 'Table Grid'

# Header baris 0
headers_bp = ["Dimensi", "Definisi", "Aspek yang Digali"]
for i, h in enumerate(headers_bp):
    cell = tbl_bp.rows[0].cells[i]
    set_cell_bg(cell, BLUE_DARK)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = WHITE

rows_bp = [
    ("Necessities\n(Kebutuhan Utama)",
     "Kebutuhan fungsional mutlak yang harus dipenuhi sistem",
     "Proses absensi, pencatatan kehadiran, rekap data, notifikasi"),
    ("Lacks\n(Kekurangan Sistem Saat Ini)",
     "Masalah/gap pada sistem manual yang menjadi motivasi pengembangan",
     "Kesulitan rekap, rawan manipulasi, keterlambatan informasi"),
    ("Wants\n(Keinginan Pengguna)",
     "Fitur tambahan yang diharapkan pengguna (nice-to-have)",
     "Izin digital, statistik personal, ekspor laporan, notifikasi push"),
    ("Interface & Usability\n(Antarmuka & Kegunaan)",
     "Kemudahan penggunaan, tampilan, dan aksesibilitas sistem",
     "Tampilan mobile-friendly, kemudahan navigasi, ukuran font, respons sistem"),
]
for r_idx, (dim, defn, asp) in enumerate(rows_bp):
    row = tbl_bp.rows[r_idx + 1]
    bg = BLUE_LIGHT if r_idx % 2 == 0 else GRAY_LIGHT
    for c_idx, text in enumerate([dim, defn, asp]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(10)

# Lebar kolom
widths = [Cm(3.8), Cm(6.4), Cm(6.4)]
for row in tbl_bp.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# FUNGSI PEMBUATAN TABEL IDENTITAS
# ══════════════════════════════════════════════════════════

def add_identitas_table(doc, aktor_label):
    """Buat tabel identitas narasumber."""
    heading_para(doc, f"C.{aktor_label} DATA IDENTITAS NARASUMBER", size=12, color=BLUE_MID, space_before=10)

    fields = [
        ("Nama Narasumber", ""),
        ("Jabatan / Status", ""),
        ("Kelas / Mata Pelajaran yang Diampu", ""),
        ("Pengalaman Mengajar / Belajar", ""),
        ("Tanggal Wawancara", ""),
        ("Tempat Wawancara", ""),
        ("Pewawancara", ""),
    ]

    tbl = doc.add_table(rows=len(fields), cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    for r_idx, (label, _) in enumerate(fields):
        row = tbl.rows[r_idx]
        bg = GRAY_LIGHT if r_idx % 2 == 0 else WHITE

        # Kolom label
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        run0 = p0.add_run(label)
        run0.font.size = Pt(10)
        run0.bold = True
        c0.width = Cm(5.5)

        # Kolom titik dua
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(":")
        run1.font.size = Pt(10)
        run1.bold = True
        c1.width = Cm(0.5)

        # Kolom isian
        c2 = row.cells[2]
        set_cell_bg(c2, WHITE)
        c2.width = Cm(10.6)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
# FUNGSI PEMBUATAN TABEL TUJUAN
# ══════════════════════════════════════════════════════════

def add_tujuan_box(doc, tujuan_list):
    heading_para(doc, "D. TUJUAN WAWANCARA", size=12, color=BLUE_MID, space_before=6, space_after=4)
    for item in tujuan_list:
        normal_para(doc, item, size=11)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
# FUNGSI PEMBUATAN TABEL PERTANYAAN
# ══════════════════════════════════════════════════════════

def add_pertanyaan_table(doc, section_label, pertanyaan_list):
    """
    pertanyaan_list: list of dict {
        'no': int,
        'dimensi': str,
        'pertanyaan': str,
        'probing': str,   # opsional
        'jawaban': str,   # jawaban ringkas
    }
    """
    heading_para(doc, f"E.{section_label} DAFTAR PERTANYAAN WAWANCARA", size=12, color=BLUE_MID, space_before=10, space_after=4)

    # Header kolom
    headers = ["No.", "Dimensi", "Pertanyaan Utama & Probing", "Jawaban / Catatan"]
    col_widths = [Cm(1.2), Cm(3.0), Cm(8.5), Cm(4.0)]

    tbl = doc.add_table(rows=1 + len(pertanyaan_list), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # Baris header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, BLUE_MID)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        cell.width = col_widths[i]

    for r_idx, item in enumerate(pertanyaan_list):
        row = tbl.rows[r_idx + 1]
        bg = BLUE_LIGHT if r_idx % 2 == 0 else WHITE

        # No
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(item['no']))
        r0.font.size = Pt(10)
        r0.bold = True
        c0.width = col_widths[0]

        # Dimensi
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(item['dimensi'])
        r1.font.size = Pt(9.5)
        r1.bold = True
        r1.font.color.rgb = BLUE_MID
        c1.width = col_widths[1]

        # Pertanyaan + Probing
        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r2 = p2.add_run(item['pertanyaan'])
        r2.font.size = Pt(10)
        if item.get('probing'):
            p2b = c2.add_paragraph()
            p2b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r2b = p2b.add_run("→ Probing: " + item['probing'])
            r2b.font.size = Pt(9)
            r2b.italic = True
            r2b.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        c2.width = col_widths[2]

        # Jawaban
        c3 = row.cells[3]
        set_cell_bg(c3, WHITE)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r3 = p3.add_run(item.get('jawaban', ''))
        r3.font.size = Pt(9.5)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        c3.width = col_widths[3]

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
# AKTOR 1: GURU / WALI KELAS
# ══════════════════════════════════════════════════════════

heading_para(doc, "━" * 60, size=11, color=BLUE_MID, space_before=10, space_after=2, bold=False)
heading_para(doc, "INSTRUMEN WAWANCARA — AKTOR 1: GURU / WALI KELAS", size=13, color=WHITE, space_before=0, space_after=0)

# Buat kotak biru untuk judul aktor
p_title = doc.paragraphs[-1]
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(0)

# Kotak judul dengan shading pada paragraf (alternatif: tetap pakai heading dengan warna)
# Kita reset dan buat ulang
p_title._element.getparent().remove(p_title._element)
doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

# Buat tabel 1x1 sebagai box judul aktor
def add_actor_title(doc, text, color_bg=BLUE_DARK, color_text=WHITE):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = color_text
    cell.width = Cm(16.6)
    doc.add_paragraph()
    return tbl

add_actor_title(doc, "INSTRUMEN WAWANCARA — AKTOR 1: GURU / WALI KELAS")

add_identitas_table(doc, "1.")

tujuan_guru = [
    "1. Memahami prosedur pencatatan kehadiran siswa yang berjalan saat ini secara manual.",
    "2. Mengidentifikasi hambatan dan kelemahan sistem absensi manual.",
    "3. Menggali kebutuhan fitur pada sistem SmartPresensi (QR Code, jurnal piket, rekap otomatis).",
    "4. Mengetahui preferensi antarmuka yang diharapkan guru agar sistem mudah digunakan.",
]
add_tujuan_box(doc, tujuan_guru)

pertanyaan_guru = [
    {
        "no": 1,
        "dimensi": "Necessities",
        "pertanyaan": "Bagaimana proses pencatatan kehadiran siswa yang Bapak/Ibu lakukan saat ini? Berapa lama waktu yang dibutuhkan untuk merekap satu kelas?",
        "probing": "Apakah ada kendala saat kertas absen rusak atau hilang?",
        "jawaban": "Guru mengisi jurnal kelas secara manual; rekap mingguan dilakukan di akhir pekan, membutuhkan ±30 menit per kelas. Kendala: tulisan tidak terbaca, kertas hilang.",
    },
    {
        "no": 2,
        "dimensi": "Lacks",
        "pertanyaan": "Apa kesulitan terbesar yang Bapak/Ibu hadapi dengan sistem absensi manual saat ini, terutama dalam hal akurasi data dan waktu pengerjaan?",
        "probing": "Pernahkah terjadi kesalahan data kehadiran yang berdampak pada siswa?",
        "jawaban": "Rekap rentan kesalahan hitung; pernah terjadi siswa dinyatakan absen padahal hadir karena tanda tangan tidak terbaca. Proses memakan waktu banyak.",
    },
    {
        "no": 3,
        "dimensi": "Necessities",
        "pertanyaan": "Fitur apa yang Bapak/Ibu anggap paling penting dalam sistem absensi digital, misalnya scan QR Code, jurnal piket otomatis, atau rekap per periode?",
        "probing": "Jika ada QR Code, apakah Bapak/Ibu lebih suka QR dicetak atau ditampilkan di layar?",
        "jawaban": "QR Code dianggap prioritas utama; guru menginginkan QR statis per kelas yang ditempel di dinding. Rekap otomatis per minggu/bulan juga sangat dibutuhkan.",
    },
    {
        "no": 4,
        "dimensi": "Wants",
        "pertanyaan": "Selain fitur dasar absensi, fitur tambahan apa yang Bapak/Ibu inginkan? Misalnya ekspor laporan ke Excel/PDF, grafik kehadiran, atau pengingat otomatis?",
        "probing": "Seberapa sering Bapak/Ibu perlu mencetak laporan kehadiran?",
        "jawaban": "Guru menginginkan ekspor ke Excel dan PDF untuk keperluan administrasi. Grafik kehadiran per kelas dianggap berguna untuk bahan laporan kepada kepala sekolah.",
    },
    {
        "no": 5,
        "dimensi": "Necessities",
        "pertanyaan": "Bagaimana Bapak/Ibu saat ini mengelola izin atau keterangan sakit siswa? Apakah ada prosedur khusus yang ingin didigitalisasi?",
        "probing": "Siapa yang berwenang menyetujui izin? Apakah cukup guru kelas atau perlu BK?",
        "jawaban": "Izin saat ini melalui surat fisik yang dikumpulkan ke wali kelas; guru mengharapkan fitur input keterangan langsung di sistem dengan persetujuan berjenjang.",
    },
    {
        "no": 6,
        "dimensi": "Interface & Usability",
        "pertanyaan": "Perangkat apa yang biasa Bapak/Ibu gunakan di sekolah (smartphone/laptop/tablet)? Tampilan seperti apa yang diharapkan agar mudah dioperasikan?",
        "probing": "Apakah Bapak/Ibu nyaman menggunakan aplikasi berbasis web di ponsel?",
        "jawaban": "Mayoritas menggunakan smartphone Android. Menginginkan tampilan sederhana, tombol besar, dan tidak perlu banyak langkah untuk mencatat kehadiran.",
    },
    {
        "no": 7,
        "dimensi": "Wants",
        "pertanyaan": "Apakah Bapak/Ibu ingin mendapatkan notifikasi otomatis, misalnya ketika siswa absen tanpa keterangan lebih dari tiga kali berturut-turut?",
        "probing": "Notifikasi via apa yang paling nyaman: SMS, WhatsApp, atau notifikasi aplikasi?",
        "jawaban": "Guru menginginkan notifikasi via WhatsApp karena lebih familier. Ambang batas 3 kali berturut-turut dianggap wajar untuk tindak lanjut.",
    },
    {
        "no": 8,
        "dimensi": "Lacks",
        "pertanyaan": "Sejauh mana Bapak/Ibu dilibatkan dalam pelaporan kehadiran kepada orang tua? Apa kendala komunikasi yang sering terjadi saat ini?",
        "probing": "Apakah orang tua sering menanyakan kehadiran anaknya langsung kepada Bapak/Ibu?",
        "jawaban": "Komunikasi dilakukan via telepon/WhatsApp pribadi yang dianggap kurang resmi. Guru mengharapkan sistem yang secara otomatis menginformasikan orang tua.",
    },
]

add_pertanyaan_table(doc, "1.", pertanyaan_guru)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# AKTOR 2: SISWA
# ══════════════════════════════════════════════════════════

add_actor_title(doc, "INSTRUMEN WAWANCARA — AKTOR 2: SISWA")

# Sesuaikan label identitas siswa
heading_para(doc, "C.2. DATA IDENTITAS NARASUMBER", size=12, color=BLUE_MID, space_before=10)

fields_siswa = [
    ("Nama Narasumber", ""),
    ("Kelas / Jurusan", ""),
    ("Jenis Kelamin", ""),
    ("Tanggal Wawancara", ""),
    ("Tempat Wawancara", ""),
    ("Pewawancara", ""),
]
tbl_s = doc.add_table(rows=len(fields_siswa), cols=3)
tbl_s.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_s.style = 'Table Grid'
for r_idx, (label, _) in enumerate(fields_siswa):
    row = tbl_s.rows[r_idx]
    bg = GRAY_LIGHT if r_idx % 2 == 0 else WHITE
    c0 = row.cells[0]; set_cell_bg(c0, bg); c0.width = Cm(5.5)
    p0 = c0.paragraphs[0]; r0 = p0.add_run(label); r0.font.size = Pt(10); r0.bold = True
    c1 = row.cells[1]; set_cell_bg(c1, bg); c1.width = Cm(0.5)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(":"); r1.font.size = Pt(10); r1.bold = True
    c2 = row.cells[2]; set_cell_bg(c2, WHITE); c2.width = Cm(10.6)
doc.add_paragraph()

tujuan_siswa = [
    "1. Memahami pengalaman siswa dalam proses absensi manual sehari-hari.",
    "2. Mengidentifikasi kesulitan yang dirasakan siswa terkait pencatatan kehadiran.",
    "3. Menggali preferensi siswa terhadap fitur scan QR Code dan rekap kehadiran pribadi.",
    "4. Mengetahui harapan siswa terhadap tampilan sistem yang mobile-friendly.",
]
add_tujuan_box(doc, tujuan_siswa)

pertanyaan_siswa = [
    {
        "no": 1,
        "dimensi": "Lacks",
        "pertanyaan": "Ceritakan pengalamanmu saat proses absen dilakukan secara manual setiap hari. Apa bagian yang paling tidak nyaman atau membuang waktu?",
        "probing": "Pernahkah kamu merasa tidak adil karena data absenmu salah tercatat?",
        "jawaban": "Siswa merasa proses tanda tangan lambat, terutama di kelas besar. Beberapa pernah tidak ikut tanda tangan karena terlambat masuk dan langsung dianggap absen.",
    },
    {
        "no": 2,
        "dimensi": "Necessities",
        "pertanyaan": "Jika ada sistem absensi menggunakan QR Code yang discan dengan ponselmu, apakah kamu tertarik menggunakannya? Apa kekhawatiranmu?",
        "probing": "Apakah kamu punya smartphone yang mendukung kamera untuk scan QR?",
        "jawaban": "Hampir semua siswa memiliki smartphone. Siswa tertarik dengan QR Code namun khawatir sinyal internet di kelas yang tidak stabil.",
    },
    {
        "no": 3,
        "dimensi": "Wants",
        "pertanyaan": "Apakah kamu ingin bisa melihat rekap kehadiranmu sendiri secara online kapan saja? Informasi apa saja yang menurutmu penting ditampilkan?",
        "probing": "Apakah kamu ingin tahu persentase kehadiranmu secara real-time?",
        "jawaban": "Siswa sangat menginginkan fitur rekap pribadi; informasi yang diharapkan: total hadir, izin, sakit, tanpa keterangan, dan persentase kehadiran untuk antisipasi tidak naik kelas.",
    },
    {
        "no": 4,
        "dimensi": "Interface & Usability",
        "pertanyaan": "Bagaimana tampilan aplikasi yang menurutmu nyaman digunakan di ponsel? Apakah kamu lebih suka tampilan sederhana atau dengan banyak fitur sekaligus?",
        "probing": "Berapa lama waktu yang masih bisa kamu toleransi untuk loading aplikasi?",
        "jawaban": "Siswa menyukai tampilan bersih dan minimalis. Loading lebih dari 5 detik dianggap terlalu lambat. Warna yang disukai: biru, putih, atau hijau.",
    },
    {
        "no": 5,
        "dimensi": "Lacks",
        "pertanyaan": "Pernahkah kamu mengalami masalah terkait surat izin atau keterangan sakit? Bagaimana proses penyerahannya sekarang dan apa kesulitannya?",
        "probing": "Apakah orang tuamu pernah kesulitan mengantarkan surat izin ke sekolah?",
        "jawaban": "Beberapa siswa pernah kehilangan poin kehadiran karena surat izin lambat diserahkan. Orang tua yang bekerja kesulitan mengantarkan surat di hari yang sama.",
    },
    {
        "no": 6,
        "dimensi": "Wants",
        "pertanyaan": "Apakah kamu ingin mendapat notifikasi atau pengingat dari sistem, misalnya pengingat jika kehadiranmu mendekati batas minimum?",
        "probing": "Mau notifikasi via aplikasi, SMS, atau WhatsApp?",
        "jawaban": "Siswa menginginkan notifikasi via WhatsApp yang langsung ke nomor pribadi. Pengingat batas minimum kehadiran (75%) dianggap sangat berguna.",
    },
    {
        "no": 7,
        "dimensi": "Necessities",
        "pertanyaan": "Menurut kamu, apakah penting data kehadiranmu langsung terlihat oleh orang tuamu? Bagaimana sebaiknya informasi itu disampaikan?",
        "probing": "Apakah kamu keberatan jika orang tua bisa melihat data absensi real-time?",
        "jawaban": "Sebagian besar siswa setuju orang tua mendapat akses; ada beberapa yang khawatir privasi jika alasan ketidakhadiran terlihat secara detail.",
    },
    {
        "no": 8,
        "dimensi": "Interface & Usability",
        "pertanyaan": "Jika sistem ini tersedia sebagai aplikasi web di browser ponsel, apakah itu sudah cukup atau kamu lebih memilih aplikasi yang bisa diinstal (APK)?",
        "probing": "Apakah memori HP-mu cukup untuk menginstal aplikasi tambahan?",
        "jawaban": "Mayoritas siswa menyatakan web browser sudah cukup asalkan responsif. Beberapa menginginkan APK jika tidak perlu login berulang.",
    },
]

add_pertanyaan_table(doc, "2.", pertanyaan_siswa)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# AKTOR 3: ORANG TUA / WALI MURID
# ══════════════════════════════════════════════════════════

add_actor_title(doc, "INSTRUMEN WAWANCARA — AKTOR 3: ORANG TUA / WALI MURID")

heading_para(doc, "C.3. DATA IDENTITAS NARASUMBER", size=12, color=BLUE_MID, space_before=10)

fields_ortu = [
    ("Nama Narasumber", ""),
    ("Nama Anak / Kelas", ""),
    ("Pekerjaan", ""),
    ("Kepemilikan Smartphone", ""),
    ("Tanggal Wawancara", ""),
    ("Tempat Wawancara", ""),
    ("Pewawancara", ""),
]
tbl_o = doc.add_table(rows=len(fields_ortu), cols=3)
tbl_o.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_o.style = 'Table Grid'
for r_idx, (label, _) in enumerate(fields_ortu):
    row = tbl_o.rows[r_idx]
    bg = GRAY_LIGHT if r_idx % 2 == 0 else WHITE
    c0 = row.cells[0]; set_cell_bg(c0, bg); c0.width = Cm(5.5)
    p0 = c0.paragraphs[0]; r0 = p0.add_run(label); r0.font.size = Pt(10); r0.bold = True
    c1 = row.cells[1]; set_cell_bg(c1, bg); c1.width = Cm(0.5)
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(":"); r1.font.size = Pt(10); r1.bold = True
    c2 = row.cells[2]; set_cell_bg(c2, WHITE); c2.width = Cm(10.6)
doc.add_paragraph()

tujuan_ortu = [
    "1. Memahami sejauh mana orang tua mengetahui kehadiran anaknya di sekolah saat ini.",
    "2. Mengidentifikasi hambatan komunikasi antara orang tua dan sekolah terkait kehadiran.",
    "3. Menggali kebutuhan orang tua terhadap fitur izin digital dan notifikasi kehadiran.",
    "4. Mengetahui kemudahan akses yang diharapkan orang tua terhadap sistem SmartPresensi.",
]
add_tujuan_box(doc, tujuan_ortu)

pertanyaan_ortu = [
    {
        "no": 1,
        "dimensi": "Lacks",
        "pertanyaan": "Bagaimana cara Bapak/Ibu mengetahui kehadiran anak di sekolah saat ini? Seberapa sering mendapatkan informasi tersebut?",
        "probing": "Apakah Bapak/Ibu pernah terkejut mengetahui anak tidak masuk sekolah padahal sudah pergi dari rumah?",
        "jawaban": "Orang tua biasanya baru mengetahui ketidakhadiran anak dari rapor akhir semester atau dari cerita anak sendiri. Beberapa pernah mengalami anak bolos tanpa sepengetahuan mereka.",
    },
    {
        "no": 2,
        "dimensi": "Necessities",
        "pertanyaan": "Seberapa penting bagi Bapak/Ibu untuk mengetahui status kehadiran anak secara real-time setiap hari? Informasi apa yang ingin diketahui?",
        "probing": "Apakah Bapak/Ibu lebih suka laporan harian, mingguan, atau segera setelah kejadian?",
        "jawaban": "Orang tua sangat menginginkan informasi real-time terutama jika anak tidak hadir tanpa keterangan. Laporan harian dianggap ideal; mingguan sebagai rekap.",
    },
    {
        "no": 3,
        "dimensi": "Wants",
        "pertanyaan": "Apakah Bapak/Ibu ingin mendapat notifikasi otomatis ke ponsel ketika anak tidak hadir atau terlambat? Melalui media apa yang paling mudah diakses?",
        "probing": "Apakah nomor WhatsApp Bapak/Ibu aktif dan sering digunakan?",
        "jawaban": "Seluruh narasumber memiliki WhatsApp aktif; notifikasi via WhatsApp dianggap paling mudah. Notifikasi diharapkan muncul dalam waktu kurang dari 15 menit setelah kejadian.",
    },
    {
        "no": 4,
        "dimensi": "Necessities",
        "pertanyaan": "Bagaimana prosedur izin anak saat ini ketika sakit atau ada keperluan keluarga? Apa kesulitan yang sering dihadapi?",
        "probing": "Pernahkah izin anak tidak tercatat di sekolah karena surat terlambat sampai?",
        "jawaban": "Proses izin melalui surat fisik atau telepon ke wali kelas; sering terjadi miskomunikasi karena guru tidak selalu bisa dihubungi. Orang tua menginginkan izin via aplikasi.",
    },
    {
        "no": 5,
        "dimensi": "Wants",
        "pertanyaan": "Apakah Bapak/Ibu tertarik dengan fitur izin digital di mana Bapak/Ibu bisa mengirimkan permohonan izin langsung dari ponsel dan mendapat konfirmasi dari sekolah?",
        "probing": "Apakah Bapak/Ibu nyaman melakukan konfirmasi/tanda tangan digital?",
        "jawaban": "Orang tua antusias dengan izin digital; sebagian perlu pendampingan awal untuk fitur tanda tangan digital. Konfirmasi via tombol persetujuan dianggap lebih mudah.",
    },
    {
        "no": 6,
        "dimensi": "Necessities",
        "pertanyaan": "Apakah Bapak/Ibu ingin bisa melihat rekap kehadiran anak secara keseluruhan, misalnya berapa kali hadir, sakit, izin, dan tanpa keterangan dalam satu semester?",
        "probing": "Apakah informasi ini berguna untuk memotivasi anak agar rajin hadir?",
        "jawaban": "Orang tua sangat menginginkan rekap semester; informasi ini digunakan untuk memantau dan memotivasi anak. Tampilan berupa grafik batang sederhana dianggap ideal.",
    },
    {
        "no": 7,
        "dimensi": "Interface & Usability",
        "pertanyaan": "Seberapa sering Bapak/Ibu menggunakan aplikasi berbasis web di ponsel? Apa kendala teknis yang biasanya dialami?",
        "probing": "Apakah Bapak/Ibu terbiasa menggunakan aplikasi seperti m-Banking atau media sosial?",
        "jawaban": "Sebagian besar familiar dengan WhatsApp dan media sosial. Kendala: lupa password dan koneksi internet tidak stabil. Menginginkan login satu langkah atau via nomor HP.",
    },
    {
        "no": 8,
        "dimensi": "Interface & Usability",
        "pertanyaan": "Tampilan seperti apa yang Bapak/Ibu harapkan dari sistem ini agar mudah dipahami meski tidak sering menggunakan teknologi? Fitur apa yang harus paling mudah ditemukan?",
        "probing": "Apakah Bapak/Ibu nyaman dengan teks kecil atau perlu tulisan yang lebih besar?",
        "jawaban": "Orang tua menginginkan tampilan bersih dengan font besar, ikon jelas, dan menu utama maksimal 4 item. Fitur 'Status Kehadiran Hari Ini' harus langsung terlihat di halaman utama.",
    },
]

add_pertanyaan_table(doc, "3.", pertanyaan_ortu)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# PENUTUP
# ══════════════════════════════════════════════════════════

add_actor_title(doc, "PENUTUP & UCAPAN TERIMA KASIH", color_bg=BLUE_MID)

normal_para(doc,
    "Pewawancara menyampaikan terima kasih yang sebesar-besarnya kepada narasumber atas "
    "kesediaan meluangkan waktu dan memberikan informasi yang sangat berharga bagi "
    "penelitian ini. Seluruh data yang diperoleh akan digunakan semata-mata untuk kepentingan "
    "akademik dan pengembangan sistem SmartPresensi demi kemajuan pelayanan pendidikan "
    "di SMKN 7 Purworejo.",
    size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6, space_after=10)

add_divider(doc, '255F A8')

# Tabel tanda tangan
heading_para(doc, "Tanda Tangan", size=11, color=BLUE_DARK, space_before=10, space_after=6)

tbl_ttd = doc.add_table(rows=4, cols=2)
tbl_ttd.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_ttd.style = 'Table Grid'

labels_ttd = [
    ("Narasumber", "Pewawancara"),
    ("", ""),
    ("", ""),
    ("(____________________)", "(____________________)")
]

for r_idx, (l, r) in enumerate(labels_ttd):
    row = tbl_ttd.rows[r_idx]
    for c_idx, text in enumerate([l, r]):
        cell = row.cells[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)
        if r_idx == 0:
            run.bold = True
            run.font.color.rgb = BLUE_DARK
        cell.width = Cm(8.3)
        # Hapus border semua kecuali baris 0
        if r_idx == 0:
            set_cell_bg(cell, BLUE_LIGHT)

# ──────────────────────────────────────────────────────────
# SIMPAN
# ──────────────────────────────────────────────────────────
output_path = r"d:\TA\Presensi\presensi-app\scratch\Instrumen_Wawancara_SmartPresensi.docx"
doc.save(output_path)
print(f"[OK] File berhasil disimpan: {output_path}")
